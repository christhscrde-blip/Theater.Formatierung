from __future__ import annotations

import os
import tempfile
from pathlib import Path

from docx import Document

from ..document_model_types import DocumentModel
from ..formatting_plan_models import FormattingPlan
from ..validation import InvalidFormattingPlan, assert_valid_formatting_plan
from ..verifier import assert_visible_text_unchanged, visible_text_hash
from .paragraph_renderer import ParagraphRenderer
from .render_context import RenderContext, RendererError


class WordRenderer:
    def __init__(self, paragraph_renderer: ParagraphRenderer | None = None):
        self._paragraph_renderer = paragraph_renderer or ParagraphRenderer()

    def render(
        self,
        input_docx: str | Path,
        output_docx: str | Path,
        document_model: DocumentModel,
        formatting_plan: FormattingPlan,
    ) -> Path:
        input_path = Path(input_docx)
        output_path = Path(output_docx)
        _assert_plan_is_valid_for_rendering(formatting_plan)
        _assert_inputs_are_consistent(input_path, document_model, formatting_plan)

        docx = Document(str(input_path))
        context = RenderContext(
            docx=docx, document_model=document_model, formatting_plan=formatting_plan
        )
        self._render_document(context)
        _save_verified_docx(context, input_path, output_path)
        return output_path

    def _render_document(self, context: RenderContext) -> None:
        if len(context.docx.paragraphs) != context.formatting_plan.paragraph_count:
            raise RendererError("DOCX paragraph count does not match FormattingPlan")
        for paragraph, paragraph_plan in zip(
            context.docx.paragraphs, context.formatting_plan.paragraphs, strict=True
        ):
            self._paragraph_renderer.render(paragraph, paragraph_plan)


def _assert_plan_is_valid_for_rendering(formatting_plan: FormattingPlan) -> None:
    try:
        assert_valid_formatting_plan(formatting_plan)
    except InvalidFormattingPlan as exc:
        raise RendererError(str(exc)) from exc


def _assert_inputs_are_consistent(
    input_path: Path, document_model: DocumentModel, formatting_plan: FormattingPlan
) -> None:
    if not document_model.has_integrity:
        raise RendererError("DocumentModel integrity check failed")
    if formatting_plan.paragraph_count != document_model.paragraph_count:
        raise RendererError(
            "FormattingPlan paragraph count does not match DocumentModel"
        )
    if not formatting_plan.has_integrity:
        raise RendererError("FormattingPlan integrity check failed")
    if formatting_plan.visible_text != document_model.visible_text:
        raise RendererError("FormattingPlan text does not match DocumentModel")
    if visible_text_hash(input_path) != document_model.visible_text_sha256:
        raise RendererError("Input DOCX visible text does not match DocumentModel hash")


def _save_verified_docx(
    context: RenderContext, input_path: Path, output_path: Path
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    tmp_path = _temporary_docx_path(output_path)
    try:
        context.docx.save(str(tmp_path))
        assert_visible_text_unchanged(input_path, tmp_path)
        os.replace(tmp_path, output_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink()


def _temporary_docx_path(output_path: Path) -> Path:
    handle = tempfile.NamedTemporaryFile(
        suffix=".docx",
        prefix=f".{output_path.stem}-",
        dir=output_path.parent,
        delete=False,
    )
    handle.close()
    return Path(handle.name)

from pathlib import Path

from docx import Document

from src.document_model import (
    build_document_model_from_docx,
    build_document_model_from_texts,
)
from src.models import ParagraphType, TextSpanType


def make_docx(path: Path, paragraphs: list[str]) -> None:
    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


def test_document_model_preserves_visible_text_and_hash_from_texts():
    texts = [
        "FIGUR A.",
        "Aber ist Euch auch wohl, Vater? (sieht ihn an)",
        "(Geht ab.)",
    ]

    model = build_document_model_from_texts(texts, source_file="memory")

    assert model.source_file == "memory"
    assert model.paragraph_count == 3
    assert model.visible_text == "\n".join(texts)
    assert model.visible_text_sha256
    assert [paragraph.text for paragraph in model.paragraphs] == texts
    assert all(paragraph.has_integrity for paragraph in model.paragraphs)


def test_document_model_stores_exactly_one_classification_per_paragraph():
    model = build_document_model_from_texts(["FIGUR A.", "Text ohne Sprecherkontext."])

    assert model.paragraphs[0].classification.type == ParagraphType.SPEAKER
    assert model.paragraphs[1].classification.type == ParagraphType.REPLIQUE
    assert model.paragraphs[1].classification.speaker == "Figur A"


def test_document_model_splits_inline_stage_without_changing_text():
    text = "Aber ist Euch wohl? (nimmt den Brief) Antwortet!"

    model = build_document_model_from_texts(["FIGUR A.", text])
    paragraph = model.paragraphs[1]

    assert paragraph.classification.type == ParagraphType.REPLIQUE
    assert [span.type for span in paragraph.spans] == [
        TextSpanType.REPLIQUE,
        TextSpanType.INLINE_STAGE,
        TextSpanType.REPLIQUE,
    ]
    assert "".join(span.text for span in paragraph.spans) == text


def test_document_model_marks_unclassified_paragraphs_for_review():
    model = build_document_model_from_texts(["Ein unklarer Absatz ohne Kontext"])

    paragraph = model.paragraphs[0]
    assert paragraph.classification.type == ParagraphType.UNCLASSIFIED
    assert paragraph.needs_manual_review
    assert "needs_manual_review" in paragraph.classification.flags


def test_document_model_from_docx_uses_docx_visible_hash(tmp_path: Path):
    docx_path = tmp_path / "sample.docx"
    paragraphs = ["Erster Akt", "FIGUR A.", "Ich komme."]
    make_docx(docx_path, paragraphs)

    model = build_document_model_from_docx(docx_path)

    assert model.source_file == str(docx_path)
    assert model.visible_text == "\n".join(paragraphs)
    assert model.paragraph_count == len(paragraphs)
    assert all(paragraph.has_integrity for paragraph in model.paragraphs)


def test_document_model_splits_speaker_with_replique_prefix_losslessly():
    text = "FIGUR A. Nachrichten von meiner Schwester?"

    model = build_document_model_from_texts([text])
    paragraph = model.paragraphs[0]

    assert paragraph.classification.type == ParagraphType.SPEAKER_WITH_REPLIQUE
    assert [span.type for span in paragraph.spans] == [
        TextSpanType.SPEAKER,
        TextSpanType.REPLIQUE,
    ]
    assert paragraph.spans[0].text == "FIGUR A. "
    assert paragraph.spans[0].speaker == "Figur A"
    assert "".join(span.text for span in paragraph.spans) == text


def test_document_model_splits_speaker_inline_stage_losslessly():
    text = "FIGUR A (leise)."

    model = build_document_model_from_texts([text])
    paragraph = model.paragraphs[0]

    assert paragraph.classification.type == ParagraphType.SPEAKER_WITH_STAGE
    assert [span.type for span in paragraph.spans] == [
        TextSpanType.SPEAKER,
        TextSpanType.INLINE_STAGE,
        TextSpanType.SPEAKER,
    ]
    assert "".join(span.text for span in paragraph.spans) == text

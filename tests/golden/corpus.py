from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

from src.report import FormattingReport


@dataclass(frozen=True)
class GoldenDocument:
    name: str
    coverage: tuple[str, ...]
    paragraphs: tuple[str, ...]
    page_break_after: tuple[int, ...] = ()
    unusual_indentation: bool = False

    def write_docx(self, path: Path) -> None:
        doc = Document()
        for index, text in enumerate(self.paragraphs):
            paragraph = doc.add_paragraph(text)
            if self.unusual_indentation:
                paragraph.paragraph_format.left_indent = Pt(18 + index)
                paragraph.paragraph_format.first_line_indent = Pt(-6)
            if index in self.page_break_after:
                paragraph.add_run().add_break(WD_BREAK.PAGE)
        doc.save(path)


GOLDEN_DOCUMENTS: tuple[GoldenDocument, ...] = (
    GoldenDocument(
        name="modern_layout_inline_stage",
        coverage=("modern layout", "inline stage directions"),
        paragraphs=(
            "AKT I",
            "1. Szene",
            "Ein heller Probenraum.",
            "ANNA. Wir beginnen jetzt (nimmt das Skript).",
            "BERT. Dann höre ich zu.",
        ),
    ),
    GoldenDocument(
        name="classic_layout_multiple_speakers",
        coverage=("classic layout", "multiple speakers"),
        paragraphs=(
            "Erster Akt",
            "Erste Szene",
            "Ein Zimmer im Schloss.",
            "KÖNIG.",
            "Wer spricht vor meiner Tür?",
            "DIENER.",
            "Ein Bote aus der Stadt.",
            "KÖNIGIN.",
            "So lasst ihn ein.",
        ),
    ),
    GoldenDocument(
        name="ocr_spacing_headings",
        coverage=("OCR spacing issues",),
        paragraphs=(
            "  I.    Akl  ",
            "  2 .   Szcne  ",
            "Alter Saal im Stadttheater.",
            "FIGUR A.  Wohin   geht Ihr?",
        ),
    ),
    GoldenDocument(
        name="roman_numeral_scene",
        coverage=("Roman numerals",),
        paragraphs=(
            "II. Aufzug",
            "IV. Auftritt",
            "Ein Wald.",
            "RÖMER.",
            "Wir zählen die Zeichen.",
        ),
    ),
    GoldenDocument(
        name="long_stage_direction",
        coverage=("long stage directions",),
        paragraphs=(
            "Dritter Akt",
            "1. Szene",
            "(Die Bühne bleibt eine Weile dunkel; langsam hört man Regen, "
            "dann öffnet sich die Tür, und mehrere Gestalten treten vorsichtig ein, "
            "ohne ein Wort zu sprechen.)",
            "WÄCHTER. Bleibt stehen.",
        ),
    ),
    GoldenDocument(
        name="speaker_with_inline_stage",
        coverage=("inline stage directions", "multiple speakers"),
        paragraphs=(
            "Vierte Szene",
            "MARTA (leise). Ich habe alles gehört.",
            "PAUL. Dann sag es niemandem (blickt zur Tür).",
            "MARTA. Noch nicht.",
        ),
    ),
    GoldenDocument(
        name="empty_pages_blank_paragraphs",
        coverage=("empty pages",),
        paragraphs=(
            "Erster Aufzug",
            "",
            "",
            "Seite 1",
            "Zweite Szene",
            "LEERER. Ist jemand hier?",
        ),
    ),
    GoldenDocument(
        name="explicit_page_breaks",
        coverage=("page breaks",),
        paragraphs=(
            "I. Akt",
            "1. Auftritt",
            "Eine Schenke.",
            "WIRT. Die Tür ist offen.",
            "Seite 2",
            "GAST. Dann trete ich ein.",
        ),
        page_break_after=(3,),
    ),
    GoldenDocument(
        name="unusual_indentation",
        coverage=("unusual indentation",),
        paragraphs=(
            "Zweiter Akt",
            "2. Szene",
            "Ein Zimmer mit hohem Fenster.",
            "SCHREIBER. Diese Zeile war im Manuskript stark eingerückt.",
            "BOTIN. Und diese hing am linken Rand.",
        ),
        unusual_indentation=True,
    ),
    GoldenDocument(
        name="mixed_classic_modern_ocr",
        coverage=("modern layout", "classic layout", "OCR spacing issues"),
        paragraphs=(
            "  III.    Act  ",
            "Dritter Auftritt",
            "Großer Saal im Kloster.",
            "CHOR.",
            "Wir sprechen zusammen (sehr leise).",
            "ERSTE STIMME. Und einzeln auch.",
        ),
    ),
)


def build_regression_summary(
    results: list[tuple[GoldenDocument, FormattingReport]]
) -> str:
    document_count = len(results)
    paragraph_count = sum(report.paragraph_count for _, report in results)
    manual_review_count = sum(report.manual_review_count for _, report in results)
    warning_count = sum(report.validation_warning_count for _, report in results)
    lines = [
        "Golden regression summary",
        f"documents: {document_count}",
        f"paragraphs: {paragraph_count}",
        f"manual_review: {manual_review_count}",
        f"validation_warnings: {warning_count}",
        "cases:",
    ]
    for document, report in results:
        coverage = ", ".join(document.coverage)
        lines.append(
            f"- {document.name}: {report.paragraph_count} paragraphs; {coverage}; "
            f"hash={report.output_hash[:12]}"
        )
    return "\n".join(lines)

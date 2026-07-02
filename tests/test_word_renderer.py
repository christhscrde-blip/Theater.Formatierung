from dataclasses import replace
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.document_model import build_document_model_from_docx
from src.formatting_plan import build_formatting_plan
from src.formatting_plan_models import FormattingPlan
from src.models import TextSpanType
from src.renderer.paragraph_renderer import ParagraphRenderer
from src.renderer.render_context import RendererError
from src.renderer.run_formatter import RunFormatter
from src.renderer.word_renderer import WordRenderer
from src.style_engine import StyleEngine
from src.style_loader import load_style_sheet_from_text
from src.verifier import assert_visible_text_unchanged, visible_text_hash

STYLE = """
name: renderer-test
defaults:
  font_family: Test Serif
  font_size: 12
  bold: false
  italic: false
  underline: false
  text_color: "#111111"
  highlight_color: null
  alignment: left
  left_indent: 0
  right_indent: 0
  first_line_indent: 0
  spacing_before: 0
  spacing_after: 0
  line_spacing: 1.0
styles:
  speaker:
    bold: true
    text_color: "#222222"
    alignment: center
    spacing_before: 6
    spacing_after: 3
    left_indent: 4
    right_indent: 2
    first_line_indent: 1
  replique:
    text_color: "#000000"
  inline_stage:
    italic: true
    text_color: "#777777"
    highlight_color: "#FFFF00"
"""


def make_engine() -> StyleEngine:
    return StyleEngine(load_style_sheet_from_text(STYLE))


def make_docx(path: Path, paragraphs: list[str]) -> None:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)


def make_plan(docx_path: Path) -> FormattingPlan:
    model = build_document_model_from_docx(docx_path)
    return build_formatting_plan(model, make_engine())


def test_run_formatter_applies_font_and_highlight():
    sheet = load_style_sheet_from_text(STYLE)
    doc = Document()
    run = doc.add_paragraph().add_run("(leise)")

    RunFormatter().apply(run, sheet.style_for(TextSpanType.INLINE_STAGE))

    assert run.font.name == "Test Serif"
    assert run.font.size.pt == 12
    assert run.font.italic is True
    assert str(run.font.color.rgb) == "777777"
    assert run._element.xml.find("FFFF00") >= 0


def test_paragraph_renderer_applies_paragraph_formatting(tmp_path: Path):
    docx_path = tmp_path / "paragraph.docx"
    make_docx(docx_path, ["FIGUR A."])
    model = build_document_model_from_docx(docx_path)
    plan = build_formatting_plan(model, make_engine())
    doc = Document(str(docx_path))

    ParagraphRenderer().render(doc.paragraphs[0], plan.paragraphs[0])

    paragraph = doc.paragraphs[0]
    assert paragraph.text == "FIGUR A."
    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paragraph.paragraph_format.space_before.pt == 6
    assert paragraph.paragraph_format.space_after.pt == 3
    assert paragraph.runs[0].bold is True


def test_word_renderer_rejects_invalid_formatting_plan_hash(tmp_path: Path):
    input_docx = tmp_path / "input.docx"
    make_docx(input_docx, ["FIGUR A."])
    model = build_document_model_from_docx(input_docx)
    plan = build_formatting_plan(model, make_engine())
    broken_plan = replace(plan, visible_text_sha256="0" * 64)

    with pytest.raises(RendererError, match="visible-text hash"):
        WordRenderer().render(input_docx, tmp_path / "out.docx", model, broken_plan)


def test_word_renderer_rejects_invalid_plan_before_rendering(tmp_path: Path):
    input_docx = tmp_path / "input.docx"
    make_docx(input_docx, ["FIGUR A.", "Hallo"])
    model = build_document_model_from_docx(input_docx)
    plan = build_formatting_plan(model, make_engine())
    shorter_plan = replace(plan, paragraphs=plan.paragraphs[:1])

    with pytest.raises(RendererError, match="visible-text hash"):
        WordRenderer().render(input_docx, tmp_path / "out.docx", model, shorter_plan)


def test_word_renderer_renders_docx_without_changing_visible_text(tmp_path: Path):
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    make_docx(input_docx, ["FIGUR A. Hallo (leise)"])
    model = build_document_model_from_docx(input_docx)
    plan = build_formatting_plan(model, make_engine())

    WordRenderer().render(input_docx, output_docx, model, plan)

    assert output_docx.exists()
    assert visible_text_hash(input_docx) == visible_text_hash(output_docx)
    assert_visible_text_unchanged(input_docx, output_docx)

    rendered = Document(str(output_docx))
    paragraph = rendered.paragraphs[0]
    assert paragraph.text == "FIGUR A. Hallo (leise)"
    assert [run.text for run in paragraph.runs] == ["FIGUR A. ", "Hallo ", "(leise)"]
    assert paragraph.runs[0].bold is True
    assert paragraph.runs[2].italic is True
    assert paragraph.runs[2]._element.xml.find("FFFF00") >= 0
